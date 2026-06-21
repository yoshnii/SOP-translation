#!/usr/bin/env python3
"""
docx 保格式翻译 — 内部网站后端 (FastAPI)。

功能:上传 .docx → 豆包翻译(保留格式/图片/术语表)→ 在线预览/下载,
并计算字符数与费用(供使用者结算)。

环境变量:
  ARK_API_KEY   服务器统一使用的火山方舟 key(必填)
  ARK_MODEL     模型,默认 doubao-seed-translation-250915
启动:  uvicorn app:app --host 0.0.0.0 --port 8000
"""
from __future__ import annotations

import os
import re
import threading
import uuid
import zipfile
from datetime import datetime, timezone

from fastapi import FastAPI, UploadFile, File, HTTPException
from fastapi.responses import FileResponse, HTMLResponse, JSONResponse
from fastapi.staticfiles import StaticFiles

import translator_core as core
from docx import Document

# ----------------------------------------------------------------------------
# 配置
# ----------------------------------------------------------------------------
BASE = os.path.dirname(os.path.abspath(__file__))
DATA = os.path.join(BASE, "data")
os.makedirs(DATA, exist_ok=True)

GLOSSARY_PATH = os.path.join(BASE, "glossary.csv")
OVERRIDES_PATH = os.path.join(BASE, "overrides.csv")

# 翻译引擎:DeepSeek V4-Pro(章节级,上下文消歧,质量最高)
DS_MODEL = os.environ.get("DS_MODEL", "deepseek-v4-pro-260425")
# DeepSeek V4-Pro 价格(火山方舟,元/百万 token)
DS_PRICE_IN = 4.0
DS_PRICE_OUT = 16.0

# 预估用:输出字符 ≈ 输入字符 × 该系数(中->英膨胀经验值)
OUT_RATIO = 1.8

# 内存任务表:job_id -> {status, done, total, ...}
JOBS: dict[str, dict] = {}

app = FastAPI(title="DOCX 保格式翻译")

# CORS:允许 Cloudflare Pages 前端跨域调用本地 API。
# 默认允许所有来源(内部工具,简单);如需收紧,把 allow_origins 改成你的 Pages 域名。
from fastapi.middleware.cors import CORSMiddleware
_ALLOWED = os.environ.get("ALLOWED_ORIGINS", "*")
app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"] if _ALLOWED == "*" else [o.strip() for o in _ALLOWED.split(",")],
    allow_methods=["*"],
    allow_headers=["*"],
)


# ----------------------------------------------------------------------------
# 术语表 API
# ----------------------------------------------------------------------------
import csv as _csv

# 添加术语的密码(改这里,或用环境变量 GLOSSARY_PASSWORD)
GLOSSARY_PASSWORD = os.environ.get("GLOSSARY_PASSWORD", "bgi2026")


def _read_glossary_rows():
    """直接读 CSV,返回 [(zh, en, genre)],跳过注释/表头。"""
    items = []
    if not os.path.exists(GLOSSARY_PATH):
        return items
    for row in _csv.reader(open(GLOSSARY_PATH, newline="", encoding="utf-8")):
        if not row or row[0].strip().startswith("#") or len(row) < 2:
            continue
        zh, en = row[0].strip(), row[1].strip()
        if not (zh and en and zh != "中文"):
            continue
        genre = row[4].strip() if len(row) >= 5 and row[4].strip() else "sci"
        items.append({"zh": zh, "en": en, "genre": genre})
    return items


@app.get("/api/glossary")
def get_glossary():
    """返回术语表(含分类),供前端展示。"""
    items = _read_glossary_rows()
    return {"count": len(items), "items": items}


@app.post("/api/glossary/add")
def add_glossary(body: dict):
    """添加一条术语(需密码)。body: {zh, en, genre, password}"""
    if (body or {}).get("password", "") != GLOSSARY_PASSWORD:
        raise HTTPException(403, "密码错误")
    zh = (body.get("zh") or "").strip()
    en = (body.get("en") or "").strip()
    genre = (body.get("genre") or "sci").strip()
    if not zh or not en:
        raise HTTPException(400, "中文和英文都要填")
    if genre not in ("sci", "siro"):
        genre = "sci"
    # 查重
    for it in _read_glossary_rows():
        if it["zh"] == zh:
            raise HTTPException(400, f"术语「{zh}」已存在(译法:{it['en']})")
    # 追加到 CSV(中文,英文,备注,错译,分类)
    with open(GLOSSARY_PATH, "a", newline="", encoding="utf-8") as f:
        _csv.writer(f).writerow([zh, en, "[网页添加]", "", genre])
    return {"ok": True, "zh": zh, "en": en, "genre": genre}


# ----------------------------------------------------------------------------
# 上传 + 翻译
# ----------------------------------------------------------------------------
@app.post("/api/upload")
async def upload(file: UploadFile = File(...)):
    """接收 .docx,先做统计(段数/字符/预估费用),返回 job_id。不立即翻译。"""
    if not file.filename.lower().endswith(".docx"):
        raise HTTPException(400, "只支持 .docx 文件")

    job_id = uuid.uuid4().hex[:12]
    job_dir = os.path.join(DATA, job_id)
    os.makedirs(job_dir, exist_ok=True)
    in_path = os.path.join(job_dir, "input.docx")
    with open(in_path, "wb") as f:
        f.write(await file.read())

    # 统计
    try:
        doc = Document(in_path)
    except Exception:
        raise HTTPException(400, "无法读取该 .docx,文件可能损坏")
    segments = core.collect_segments(doc)
    to_translate = [s for s in segments if not core.should_skip(s.text)]
    unique = list(dict.fromkeys(s.text for s in to_translate))
    in_chars = sum(len(re.sub(r"\s+", " ", t).strip()) for t in unique)
    out_chars = int(in_chars * OUT_RATIO)
    # DeepSeek V4-Pro 预估费用(按 token)
    _in_tok = int(in_chars * 0.7) + 1500 * max(1, len(unique) // 40)
    _out_tok = int(out_chars * 0.35 * 2.5)
    cost = _in_tok / 1e6 * DS_PRICE_IN + _out_tok / 1e6 * DS_PRICE_OUT

    with zipfile.ZipFile(in_path) as z:
        n_img = len([n for n in z.namelist()
                     if n.startswith("word/media/") and "." in n.split("/")[-1]])

    JOBS[job_id] = {
        "status": "ready",         # ready -> running -> done / error
        "filename": file.filename,
        "dir": job_dir,
        "in_path": in_path,
        "total_segments": len(segments),
        "translate_segments": len(to_translate),
        "unique_segments": len(unique),
        "images": n_img,
        "in_chars": in_chars,
        "out_chars_est": out_chars,
        "cost_est": round(cost, 3),
        "done": 0,
        "total": len(unique),
        "created": datetime.now(timezone.utc).isoformat(),
    }
    return JOBS_public(job_id)


@app.post("/api/translate/{job_id}")
def start_translate(job_id: str, body: dict = None):
    """启动翻译(后台线程)。

    body 可选 {"api_key": "ark-..."}:用户自填 key 则用它(各付各的,
    不显示乞讨);否则用服务器默认 key(走我的额度,翻完显示乞讨求转账)。
    """
    job = JOBS.get(job_id)
    if not job:
        raise HTTPException(404, "任务不存在")
    if job["status"] == "running":
        return JOBS_public(job_id)

    user_key = (body or {}).get("api_key", "").strip() if body else ""
    if user_key:
        job["api_key"] = user_key
        job["own_key"] = True          # 用户自带 key
    elif os.environ.get("ARK_API_KEY"):
        job["api_key"] = os.environ["ARK_API_KEY"]
        job["own_key"] = False         # 走服务器(我的)key → 显示乞讨
    else:
        raise HTTPException(400, "请填写你的火山方舟 API Key(服务器未配置默认 key)")

    job["status"] = "running"
    job["done"] = 0
    threading.Thread(target=_run_translation, args=(job_id,), daemon=True).start()
    return JOBS_public(job_id)


def _run_translation(job_id: str):
    """后台线程:用 DeepSeek 章节级引擎翻译,更新进度。"""
    import deepseek_translator as ds
    job = JOBS[job_id]
    try:
        os.environ["ARK_API_KEY"] = job["api_key"]  # 引擎从环境读 key

        def progress(done, total):
            job["done"] = done
            job["total"] = total

        out_path = os.path.join(job["dir"], "output_en.docx")
        # DeepSeek 章节级翻译(全覆盖+保格式+术语注入+残留补翻),返回字符统计
        stats = ds.translate_doc(
            job["in_path"], out_path, DS_MODEL,
            progress_cb=progress, glossary_path=GLOSSARY_PATH,
            return_stats=True,
        )
        job["out_path"] = out_path

        # 费用:DeepSeek 按 token(火山 V4-Pro 输入4/输出16 元每百万token)
        in_chars = stats.get("in_chars", 0)
        out_chars = stats.get("out_chars", 0)
        in_tok = int(in_chars * 0.7) + 1500 * stats.get("chunks", 1)
        out_tok = int(out_chars * 0.35 * 2.5)  # V4-Pro 推理输出偏多
        cost = in_tok / 1e6 * DS_PRICE_IN + out_tok / 1e6 * DS_PRICE_OUT
        job["in_chars"] = in_chars
        job["out_chars"] = out_chars
        job["cost"] = round(cost, 2)
        job["residual_zh"] = stats.get("residual", 0)
        job["status"] = "done"
    except Exception as e:
        import traceback
        job["status"] = "error"
        job["error"] = f"{type(e).__name__}: {e}"
        print(traceback.format_exc())


@app.get("/api/status/{job_id}")
def status(job_id: str):
    if job_id not in JOBS:
        raise HTTPException(404, "任务不存在")
    return JOBS_public(job_id)


@app.get("/api/download/{job_id}")
def download(job_id: str):
    job = JOBS.get(job_id)
    if not job or job.get("status") != "done":
        raise HTTPException(404, "结果尚未就绪")
    fname = os.path.splitext(job["filename"])[0] + "_EN.docx"
    return FileResponse(job["out_path"],
                        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                        filename=fname)


@app.get("/api/preview/{job_id}")
def preview(job_id: str):
    """返回译文的纯文本预览(前若干段),供网页在线查看。"""
    job = JOBS.get(job_id)
    if not job or job.get("status") != "done":
        raise HTTPException(404, "结果尚未就绪")
    doc = Document(job["out_path"])
    paras = [p.text for p in doc.paragraphs if p.text.strip()]
    return {"paragraphs": paras[:200], "total": len(paras)}


def JOBS_public(job_id: str) -> dict:
    """对外暴露的任务信息(不含内部路径)。"""
    j = JOBS[job_id]
    pct = round(j["done"] / j["total"] * 100, 1) if j.get("total") else 0
    return {
        "job_id": job_id,
        "status": j["status"],
        "filename": j["filename"],
        "total_segments": j.get("total_segments"),
        "translate_segments": j.get("translate_segments"),
        "unique_segments": j.get("unique_segments"),
        "images": j.get("images"),
        "in_chars": j.get("in_chars"),
        "out_chars_est": j.get("out_chars_est"),
        "out_chars": j.get("out_chars"),
        "cost_est": j.get("cost_est"),
        "cost": j.get("cost"),
        "done": j.get("done"),
        "total": j.get("total"),
        "percent": pct,
        "residual_zh": j.get("residual_zh"),
        "own_key": j.get("own_key", False),
        "error": j.get("error"),
    }


# ----------------------------------------------------------------------------
# 前端静态页
# ----------------------------------------------------------------------------
@app.get("/", response_class=HTMLResponse)
def index():
    with open(os.path.join(BASE, "static", "index.html"), encoding="utf-8") as f:
        return f.read()
